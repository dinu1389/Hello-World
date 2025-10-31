#!/usr/bin/env python3
"""
Script to split Timeless_Bonds_KDP.docx into individual chapter files.
Each chapter will be saved as a separate docx file.
"""

from docx import Document
import re
import os

def split_chapters(input_file, output_dir="chapters"):
    """
    Split the document into individual chapter files.
    
    Args:
        input_file: Path to the input docx file
        output_dir: Directory to save chapter files
    """
    # Create output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Load the document
    print(f"Loading document: {input_file}")
    doc = Document(input_file)
    
    # Track chapters
    chapters = []
    current_chapter = None
    current_content = []
    
    # Chapter patterns to detect
    chapter_pattern = re.compile(r'^(Chapter\s+\d+:|#\s*Chapter\s+\d+:)\s*(.+)$', re.IGNORECASE)
    epilogue_pattern = re.compile(r'^(Epilogue:|#\s*Epilogue:)\s*(.+)$', re.IGNORECASE)
    
    print("Parsing document...")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Check if this is a chapter heading
        chapter_match = chapter_pattern.match(text)
        epilogue_match = epilogue_pattern.match(text)
        
        if chapter_match or epilogue_match:
            # Save previous chapter if exists
            if current_chapter:
                chapters.append({
                    'title': current_chapter,
                    'content': current_content
                })
            
            # Start new chapter
            if chapter_match:
                current_chapter = text
            else:
                current_chapter = text
            current_content = [para]
            print(f"Found: {current_chapter[:50]}...")
        else:
            # Add to current chapter content
            if current_chapter:
                current_content.append(para)
    
    # Don't forget the last chapter
    if current_chapter:
        chapters.append({
            'title': current_chapter,
            'content': current_content
        })
    
    print(f"\nFound {len(chapters)} chapters/sections")
    
    # Create individual files for each chapter
    for idx, chapter in enumerate(chapters, 1):
        # Create new document
        new_doc = Document()
        
        # Add title
        title_para = new_doc.add_heading(chapter['title'], level=1)
        
        # Add content
        for para in chapter['content'][1:]:  # Skip the title paragraph
            new_para = new_doc.add_paragraph()
            new_para.text = para.text
            # Copy formatting
            new_para.style = para.style
        
        # Generate filename
        # Extract chapter number or use "Epilogue"
        if 'epilogue' in chapter['title'].lower():
            filename = f"{output_dir}/Epilogue.docx"
        else:
            # Extract chapter number
            match = re.search(r'Chapter\s+(\d+)', chapter['title'], re.IGNORECASE)
            if match:
                chapter_num = match.group(1)
                # Clean title for filename
                title_clean = re.sub(r'[^\w\s-]', '', chapter['title'])
                title_clean = re.sub(r'\s+', '_', title_clean)
                filename = f"{output_dir}/Chapter_{chapter_num.zfill(2)}.docx"
            else:
                filename = f"{output_dir}/Section_{idx:02d}.docx"
        
        # Save the document
        new_doc.save(filename)
        print(f"Saved: {filename}")
    
    print(f"\nAll chapters saved to '{output_dir}/' directory")
    return len(chapters)

if __name__ == "__main__":
    input_file = "Timeless_Bonds_KDP.docx"
    
    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found!")
        exit(1)
    
    num_chapters = split_chapters(input_file)
    print(f"\n✓ Successfully split {num_chapters} chapters!")
