#!/usr/bin/env python3
"""
Script to split Timeless_Bonds_KDP.md into individual chapter docx files.
Each chapter will be saved as a separate docx file with proper formatting.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def create_formatted_docx(title, content):
    """Create a well-formatted docx document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Add title
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content paragraphs
    for line in content:
        line = line.strip()
        if not line:
            doc.add_paragraph()  # Empty paragraph for spacing
        elif line.startswith('###'):
            # Sub-sub-heading
            heading_text = line.replace('###', '').strip()
            doc.add_heading(heading_text, level=3)
        elif line.startswith('##'):
            # Sub-heading
            heading_text = line.replace('##', '').strip()
            doc.add_heading(heading_text, level=2)
        elif line.startswith('#'):
            # Main heading (shouldn't happen within chapter)
            heading_text = line.replace('#', '').strip()
            doc.add_heading(heading_text, level=1)
        else:
            # Regular paragraph
            para = doc.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    return doc

def split_chapters_from_md(input_file, output_dir="chapters"):
    """
    Split the markdown document into individual chapter docx files.
    
    Args:
        input_file: Path to the input markdown file
        output_dir: Directory to save chapter files
    """
    # Create output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Read the markdown file
    print(f"Loading markdown file: {input_file}")
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Track chapters
    chapters = []
    current_chapter = None
    current_content = []
    in_table_of_contents = False
    
    # Chapter patterns
    chapter_pattern = re.compile(r'^#\s+Chapter\s+(\d+):\s*(.+)$', re.IGNORECASE)
    epilogue_pattern = re.compile(r'^#\s+Epilogue:\s*(.+)$', re.IGNORECASE)
    toc_pattern = re.compile(r'^##\s+Table of Contents', re.IGNORECASE)
    
    print("Parsing markdown...")
    
    for line in lines:
        line = line.rstrip()
        
        # Skip table of contents section
        if toc_pattern.match(line):
            in_table_of_contents = True
            continue
        
        # Check if this is a chapter heading
        chapter_match = chapter_pattern.match(line)
        epilogue_match = epilogue_pattern.match(line)
        
        if chapter_match:
            # End table of contents if we were in it
            in_table_of_contents = False
            
            # Save previous chapter if exists
            if current_chapter:
                chapters.append({
                    'title': current_chapter,
                    'number': current_chapter_num,
                    'content': current_content
                })
            
            # Start new chapter
            current_chapter_num = chapter_match.group(1)
            chapter_title = chapter_match.group(2).strip()
            current_chapter = f"Chapter {current_chapter_num}: {chapter_title}"
            current_content = []
            print(f"Found: {current_chapter}")
            
        elif epilogue_match:
            # End table of contents if we were in it
            in_table_of_contents = False
            
            # Save previous chapter if exists
            if current_chapter:
                chapters.append({
                    'title': current_chapter,
                    'number': current_chapter_num,
                    'content': current_content
                })
            
            # Start epilogue
            epilogue_title = epilogue_match.group(1).strip()
            current_chapter = f"Epilogue: {epilogue_title}"
            current_chapter_num = "Epilogue"
            current_content = []
            print(f"Found: {current_chapter}")
            
        elif current_chapter and not in_table_of_contents:
            # Add to current chapter content (skip the chapter title line itself)
            if not (chapter_pattern.match(line) or epilogue_pattern.match(line)):
                current_content.append(line)
    
    # Don't forget the last chapter
    if current_chapter:
        chapters.append({
            'title': current_chapter,
            'number': current_chapter_num,
            'content': current_content
        })
    
    print(f"\nFound {len(chapters)} chapters/sections")
    
    # Create individual docx files for each chapter
    for chapter in chapters:
        # Create document
        doc = create_formatted_docx(chapter['title'], chapter['content'])
        
        # Generate filename
        if chapter['number'] == 'Epilogue':
            filename = f"{output_dir}/Epilogue.docx"
        else:
            chapter_num = chapter['number'].zfill(2)
            filename = f"{output_dir}/Chapter_{chapter_num}.docx"
        
        # Save the document
        doc.save(filename)
        print(f"Saved: {filename} ({len(chapter['content'])} lines)")
    
    print(f"\n✓ All chapters saved to '{output_dir}/' directory")
    return len(chapters)

if __name__ == "__main__":
    input_file = "Timeless_Bonds_KDP.md"
    
    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found!")
        exit(1)
    
    num_chapters = split_chapters_from_md(input_file)
    print(f"\n✓ Successfully split {num_chapters} chapters!")
